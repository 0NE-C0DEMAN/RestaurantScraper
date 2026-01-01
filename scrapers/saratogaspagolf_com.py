import json
import sys
import re
import requests
from pathlib import Path
from typing import Dict, List
import tempfile
import os

# Add parent directory to path to import modules
sys.path.append(str(Path(__file__).parent.parent))

try:
    from docx import Document  # pyright: ignore[reportMissingImports]
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    print("Warning: python-docx not installed.")
    print("Install with: pip install python-docx")

def scrape_saratogaspagolf_menu(url: str) -> List[Dict]:
    """Scrape menu items from saratogaspagolf.com"""
    all_items = []
    restaurant_name = "58 Roosevelt Bar & Grill"
    menu_docx_url = "https://www.saratogaspagolf.com/images/downloads/Menu_2025_cover.docx"
    
    print(f"Scraping: {url}")
    print(f"Menu document: {menu_docx_url}")
    
    # Create output directory if it doesn't exist
    output_dir = Path(__file__).parent.parent / 'output'
    output_dir.mkdir(parents=True, exist_ok=True)
    
    # Create output filename based on URL
    url_safe = url.replace('https://', '').replace('http://', '').replace('/', '_').replace('.', '_')
    # Remove 'www_' prefix and 'menu' if present
    if url_safe.startswith('www_'):
        url_safe = url_safe[4:]
    url_safe = url_safe.replace('_menu', '')
    output_json = output_dir / f'{url_safe}.json'
    print(f"Output file: {output_json}\n")
    
    if not DOCX_AVAILABLE:
        print("ERROR: python-docx is required to extract menu from Word document.")
        print("Please install it with: pip install python-docx")
        return []
    
    try:
        # Headers for downloading the document
        headers = {
            "Accept": "text/html,application/xhtml+xml,application/xml;q=0.9,image/avif,image/webp,image/apng,*/*;q=0.8,application/signed-exchange;v=b3;q=0.7",
            "Accept-Encoding": "gzip, deflate, br, zstd",
            "Accept-Language": "en-IN,en;q=0.9,hi-IN;q=0.8,hi;q=0.7,en-GB;q=0.6,en-US;q=0.5",
            "Cache-Control": "no-cache",
            "Cookie": "b690747afcde5398dcaa50974124cdb6=7f7162e20c3f5c6bca2ba56d77283210; _ga=GA1.1.1214225990.1767182442; _ga_WKM7E3SYK8=GS2.1.s1767182441$o1$g1$t1767182551$j60$l0$h0",
            "Pragma": "no-cache",
            "Priority": "u=0, i",
            "Referer": "https://www.saratogaspagolf.com/dining/restaurant",
            "Sec-Ch-Ua": '"Google Chrome";v="143", "Chromium";v="143", "Not A(Brand";v="24"',
            "Sec-Ch-Ua-Mobile": "?0",
            "Sec-Ch-Ua-Platform": '"Windows"',
            "Sec-Fetch-Dest": "document",
            "Sec-Fetch-Mode": "navigate",
            "Sec-Fetch-Site": "same-origin",
            "Sec-Fetch-User": "?1",
            "Upgrade-Insecure-Requests": "1",
            "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/143.0.0.0 Safari/537.36"
        }
        
        print("Downloading menu document...")
        response = requests.get(menu_docx_url, headers=headers, timeout=60)
        response.raise_for_status()
        
        # Save to temporary file
        with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp_file:
            tmp_file.write(response.content)
            tmp_file_path = tmp_file.name
        
        print("Extracting text from Word document...")
        # Open the document
        doc = Document(tmp_file_path)
        
        # Extract all text from the document
        full_text = []
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                full_text.append(paragraph.text.strip())
        
        # Also check tables for menu items
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text:
                        row_text.append(cell_text)
                if row_text:
                    full_text.append(" | ".join(row_text))
        
        # Clean up temp file
        os.unlink(tmp_file_path)
        
        print(f"Extracted {len(full_text)} lines from document")
        
        # Debug: Print first 30 lines to understand structure
        print("\nFirst 30 lines of extracted text:")
        for idx, line in enumerate(full_text[:30], 1):
            print(f"  {idx:2d}: {line}")
        print()
        
        print("Parsing menu items...\n")
        
        # Parse menu items from the text
        items = parse_menu_from_text(full_text)
        
        if items:
            for item in items:
                item['restaurant_name'] = restaurant_name
                item['restaurant_url'] = url
                item['menu_type'] = "Menu"  # Single menu type
            all_items.extend(items)
            print(f"[OK] Extracted {len(items)} items from menu\n")
        else:
            print(f"[WARNING] No items extracted from menu\n")
        
    except Exception as e:
        print(f"Error during scraping: {e}")
        import traceback
        traceback.print_exc()
    
    # Save JSON file with menu items
    with open(output_json, 'w', encoding='utf-8') as f:
        json.dump(all_items, f, indent=2, ensure_ascii=False)
    
    print(f"\n{'='*60}")
    print("SCRAPING COMPLETE")
    print(f"{'='*60}")
    print(f"Restaurant: {restaurant_name}")
    print(f"Total items found: {len(all_items)}")
    print(f"Saved to: {output_json}")
    
    return all_items

def parse_menu_from_text(text_lines: List[str]) -> List[Dict]:
    """Parse menu items from extracted text lines."""
    items = []
    
    # Section headers to skip
    section_headers = ['APPETIZERS', 'SALADS', 'ENTREES', 'SANDWICHES', 'BURGERS', 'DESSERTS', 
                       'BEVERAGES', 'DRINKS', 'SOUP', 'SIDES', 'KIDS', 'CHILDREN', 'FOR THE KIDS']
    
    # Skip patterns for non-menu content
    skip_patterns = [
        'NESTLED', 'HISTORIC', 'SARATOGA SPA', 'STATE PARK', 'GOLF COURSE', 'VICTORIA POOL',
        'ROOSEVELT', 'SPAC CONCERTS', 'PERFORMING ARTS', 'PHONE:', 'EMAIL:', 'LUNCH:', 
        'DAILY', 'HOURS', 'RESERVATIONS', 'PAGE', 'MENU', '2025', 'SUMMER', 'SPRING', 'FALL',
        'SERVED WITH', 'CHOICE OF', 'OPTIONS', 'REVERSE', 'NO SPLIT', 'NO HALF',
        'AGES', 'EXCEPTIONS', 'SIDE OF', 'ADD $', 'CHEESE OPTIONS', 'DRESSING OPTIONS',
        'HOMEMADE', 'DRAFTS', 'TAP SELECTION', 'ASK', 'WAIT-STAFF', 'SERVER', 'SELECTIONS',
        'CAN BEER', 'VODKA', 'GIN', 'TEQUILA', 'CANS', 'LEMONADE', 'ICED TEA', 'PEPSI',
        'SPLIT PLATE', 'CHARGE', 'DISCOUNT', 'CASH', 'PAYMENT', 'GLUTEN-FREE', 'BUN ITEMS'
    ]
    
    i = 0
    while i < len(text_lines):
        line = text_lines[i].strip()
        
        
        # Skip empty lines
        if not line:
            i += 1
            continue
        
        # Skip section headers
        line_upper = line.upper()
        if line_upper in section_headers or (line_upper.endswith('S') and line_upper[:-1] in section_headers):
            i += 1
            continue
        
        # Skip lines that are just section notes (all caps, no price)
        if line.isupper() and not re.search(r'\$?\d+\.?\d*\s*$', line) and len(line) > 20:
            i += 1
            continue
        
        # Pattern: "ITEM NAME $XX.XX" or "ITEM NAME XX.XX" or "ITEM NAME (DESC) $XX.XX"
        price_match = re.search(r'[\s$]*(\d+\.?\d*)\s*$', line)
        
        # Skip headers/footers/common text (but NOT for lines with prices - those are menu items)
        if not price_match and any(skip in line_upper for skip in skip_patterns):
            i += 1
            continue
        if price_match:
            price = price_match.group(1)
            
            # Remove price from line
            name_line = re.sub(r'[\s$]*\d+\.?\d*\s*$', '', line).strip()
            name_line = re.sub(r'^[-–]\s*', '', name_line).strip()
            
            # Extract description from parentheses in the same line
            description = ""
            name = name_line
            
            # Check for description in parentheses on same line: "ITEM (DESC) $XX"
            paren_match = re.search(r'\(([^)]+)\)', name_line)
            if paren_match:
                description = paren_match.group(1).strip()
                # Remove the parentheses part from name
                name = re.sub(r'\([^)]+\)', '', name_line).strip()
            
            # Skip if name is too short or looks like a description/note
            if len(name) < 3 or name.startswith('(') or name.upper() in section_headers:
                i += 1
                continue
            
            # Check next line(s) for description in parentheses
            desc_lines = []
            add_ons = []  # Track add-on items to include in description
            j = i + 1
            consecutive_non_desc = 0  # Track consecutive lines that aren't descriptions
            
            
            while j < len(text_lines):
                next_line = text_lines[j].strip()
                
                # Stop if we hit a section header
                if next_line.upper() in section_headers or (next_line.upper().endswith('S') and next_line.upper()[:-1] in section_headers):
                    break
                
                # Stop if we hit a skip pattern (but not for description lines in parentheses or "Served" lines)
                # Description lines in parentheses should not be skipped based on skip patterns
                # "Served" lines (like "Served w/ Homemade Au Jus on the side") should also not be skipped
                is_description_line = next_line.startswith('(') and next_line.endswith(')') and not next_line.upper().startswith('(ADD')
                is_served_line = 'SERVED' in next_line.upper() and 'SERVED WITH' not in next_line.upper()
                if not is_description_line and not is_served_line and any(skip in next_line.upper() for skip in skip_patterns):
                    break
                
                # Check if this line has a price (match $XX.XX or just XX.XX at the end for add-ons)
                price_match_next = re.search(r'\$(\d+\.?\d*)', next_line)
                # Also check for prices without $ sign (like ".75" for add-ons)
                if not price_match_next and 'ADD' in next_line.upper():
                    # Match .XX format (like ".75") or XX.XX at the end
                    price_match_next = re.search(r'([\.]\d+|\d+\.\d+)\s*$', next_line)
                
                # Check if it's an add-on item (low price, mentions "Add")
                is_add_on = False
                if price_match_next:
                    next_price = float(price_match_next.group(1))
                    # Check if this line contains multiple add-ons (with or without "/")
                    # (e.g., "Add Cheese $1.29/Add Bacon $2.06" or "Add Cheese $1.29 Add Bacon $3.86")
                    if 'ADD' in next_line.upper():
                        # Find all prices (with or without $)
                        prices_with_dollar = re.findall(r'\$(\d+\.?\d*)', next_line)
                        # Match .XX format (like ".75") or XX.XX format
                        prices_without_dollar = re.findall(r'([\.]\d+|\d+\.\d+)', next_line)
                        prices = prices_with_dollar + prices_without_dollar
                        if len(prices) > 1:
                            # Multiple add-ons - split by "/" or by "Add" pattern
                            if '/' in next_line:
                                add_on_parts = next_line.split('/')
                            else:
                                # Split by "Add" pattern (e.g., "Add Cheese $1.29 Add Bacon $3.86")
                                add_on_parts = re.split(r'(?=Add\s)', next_line, flags=re.IGNORECASE)
                            
                            for part in add_on_parts:
                                part = part.strip().strip('()')
                                if part and 'ADD' in part.upper():
                                    # Try to find price with $ first, then without
                                    part_price_match = re.search(r'\$(\d+\.?\d*)', part)
                                    if not part_price_match:
                                        # Match .XX format (like ".75") or XX.XX format
                                        part_price_match = re.search(r'([\.]\d+|\d+\.\d+)\s*$', part)
                                    if part_price_match:
                                        part_price = float(part_price_match.group(1))
                                        part_name = re.sub(r'\$\d+\.?\d*', '', part).strip()
                                        part_name = re.sub(r'\d+\.?\d*\s*$', '', part_name).strip()
                                        # Only add if it mentions "Add" and price is reasonable
                                        if part_name.upper().startswith('ADD ') and part_price < 10.0:
                                            add_ons.append(f"{part_name} ${part_price:.2f}")
                            j += 1
                            consecutive_non_desc = 0
                            continue
                    
                    # Single add-on - remove price from line to get name part
                    # Handle both $XX.XX and XX.XX formats (including .XX)
                    name_part = re.sub(r'\$\d+\.?\d*', '', next_line).strip()
                    name_part = re.sub(r'([\.]\d+|\d+\.\d+)\s*$', '', name_part).strip()
                    # Add-ons: low price (< $10) and explicitly mention "Add" (not just short names)
                    # Short names with prices >= $5 are likely menu items, not add-ons
                    # Also include items with prices < $1 (like .75) if they mention "Add"
                    if (next_price < 10.0 and 
                        (name_part.upper().startswith('ADD ') or 
                         (name_part.startswith('(') and 'ADD' in name_part.upper()) or
                         (len(name_part) < 20 and next_price < 5.0) or
                         (next_price < 1.0 and 'ADD' in name_part.upper()))):
                        is_add_on = True
                        # Collect the add-on text (remove parentheses if present, remove original price)
                        add_on_text = name_part.strip('()').strip()
                        # Remove any price that might still be in the text
                        # First remove prices with $, then remove .XX or XX.XX formats
                        add_on_text = re.sub(r'\$\d+\.?\d*', '', add_on_text).strip()
                        add_on_text = re.sub(r'\s*[\.]?\d+\.\d+\s*$', '', add_on_text).strip()
                        # Clean up any trailing dots or spaces
                        add_on_text = re.sub(r'\.\s*$', '', add_on_text).strip()
                        # Single add-on - format price with 2 decimals (always show as $X.XX)
                        add_ons.append(f"{add_on_text} ${next_price:.2f}")
                        j += 1
                        consecutive_non_desc = 0  # Reset counter when we find an add-on
                        continue
                    # Main menu item (higher price or substantial name)
                    elif next_price >= 5.0 or len(name_part) > 15:
                        # This is a new menu item, stop here
                        # j currently points to this new menu item, so we'll process it in the next iteration
                        break
                
                # Collect description lines (in parentheses or descriptive text)
                if next_line.startswith('(') and next_line.endswith(')'):
                    desc_text = next_line.strip('()').strip()
                    # Check if it's an add-on in parentheses (e.g., "(Add Turkey or Ham $4.38)" or "(ADD Pepperoni .75)")
                    # Add-ons have prices and mention "Add"
                    add_on_price_match = re.search(r'\$(\d+\.?\d*)', next_line)
                    # Also check for prices without $ sign (like ".75" for add-ons)
                    if not add_on_price_match and ('ADD' in desc_text.upper() or desc_text.upper().startswith('ADD')):
                        # Match .XX format (like ".75") or XX.XX at the end
                        add_on_price_match = re.search(r'([\.]\d+|\d+\.\d+)\s*\)?$', next_line)
                    if add_on_price_match and ('ADD' in desc_text.upper() or desc_text.upper().startswith('ADD')):
                        # This is an add-on, collect it
                        add_on_price = float(add_on_price_match.group(1))
                        # Remove the original price from the text and replace with formatted price
                        # First remove prices with $, then remove .XX or XX.XX formats (including in parentheses)
                        add_on_text = desc_text
                        # Remove prices with $ sign
                        add_on_text = re.sub(r'\$\d+\.?\d*', '', add_on_text).strip()
                        # Remove .XX or XX.XX patterns (match the exact price we found)
                        price_str = add_on_price_match.group(1)
                        # Escape the price string for regex (handle .75 and 4.38 formats)
                        price_pattern = re.escape(price_str)
                        add_on_text = re.sub(r'\s*' + price_pattern + r'\s*\)?', '', add_on_text).strip()
                        # Clean up any trailing dots, spaces, or parentheses
                        add_on_text = re.sub(r'[\.\s\)]+$', '', add_on_text).strip()
                        # Remove leading/trailing parentheses if they're now empty
                        add_on_text = add_on_text.strip('()').strip()
                        # Format price with 2 decimals (always show as $X.XX)
                        add_ons.append(f"{add_on_text} ${add_on_price:.2f}")
                        j += 1
                        consecutive_non_desc = 0  # Reset counter when we find an add-on
                        continue
                    # Regular description (not an add-on)
                    if desc_text and not desc_text.upper().startswith('ADD'):
                        desc_lines.append(desc_text)
                        consecutive_non_desc = 0  # Reset counter when we find a description
                    j += 1
                elif not next_line.isupper() and len(next_line) > 5 and not price_match_next:
                    # Descriptive text (not all caps, no price, reasonable length)
                    # Include lines like "Served w/ Homemade Au Jus on the side"
                    # Always include lines that contain "Served" (they're descriptions)
                    # But stop after collecting a "Served" line to avoid picking up section notes
                    if 'SERVED' in next_line.upper() and 'SERVED WITH' not in next_line.upper():
                        desc_lines.append(next_line)
                        consecutive_non_desc = 0  # Reset counter when we find a description
                        j += 1
                        # After a "Served" line, check if next line is a section note and stop if so
                        if j < len(text_lines):
                            next_next_line = text_lines[j].strip() if j < len(text_lines) else ""
                            if (next_next_line.isupper() and len(next_next_line) > 10) or any(skip in next_next_line.upper() for skip in ['SERVED WITH', 'CHOICE OF', 'NO HALF', 'NO SPLIT']):
                                break
                        continue
                    # Other descriptive text
                    if not any(skip in next_line.upper() for skip in ['SERVED WITH', 'CHOICE OF', 'OPTIONS ON']):
                        desc_lines.append(next_line)
                        consecutive_non_desc = 0  # Reset counter when we find a description
                    j += 1
                else:
                    # Not a description line - but check if it's a "Served" line that should be included
                    if 'SERVED' in next_line.upper() and 'SERVED WITH' not in next_line.upper():
                        # Include "Served" lines even if they don't match other patterns
                        desc_lines.append(next_line)
                        consecutive_non_desc = 0  # Reset counter when we find a description
                        j += 1
                        continue
                    # Not a description line
                    consecutive_non_desc += 1
                    if consecutive_non_desc > 1:
                        # Two consecutive non-description lines, stop
                        break
                    j += 1
            
            # Combine descriptions
            if desc_lines:
                if description:
                    description = description + " " + " ".join(desc_lines)
                else:
                    description = " ".join(desc_lines)
            
            # Add add-ons to description
            if add_ons:
                add_ons_text = ", ".join(add_ons)
                if description:
                    description = description + ". " + add_ons_text
                else:
                    description = add_ons_text
            
            # Clean up description
            description = description.strip()
            
            # Skip add-on items (they start with "Add" or have very low prices like .75)
            if name.upper().startswith('ADD ') or float(price) < 1.0:
                i += 1
                continue
            
            items.append({
                'name': name.upper(),
                'description': description,
                'price': f"${price}"
            })
            
            # Skip the description lines we processed
            # j points to the line after the last description (or the new menu item if we broke)
            # If j points to a new menu item (has price >= $5), we want to process it next
            # So we set i = j (don't increment, process it in next iteration)
            # If j points past descriptions (no price or add-on), we skip it with i = j + 1
            next_is_menu_item = False
            if j < len(text_lines):
                next_line_check = text_lines[j].strip()
                price_check = re.search(r'[\s$]*(\d+\.?\d*)\s*$', next_line_check)
                if price_check:
                    next_price_check = float(price_check.group(1))
                    name_check = re.sub(r'[\s$]*\d+\.?\d*\s*$', '', next_line_check).strip()
                    # Check if it's a main menu item (not an add-on)
                    if next_price_check >= 5.0 or (len(name_check) > 15 and not name_check.upper().startswith('ADD ')):
                        next_is_menu_item = True
            
            if next_is_menu_item:
                # j points to a new menu item, process it in next iteration
                i = j
            else:
                # j points past descriptions/add-ons, skip to next
                i = j + 1
        else:
            i += 1
    
    return items

def main():
    url = "https://www.saratogaspagolf.com/dining/restaurant"
    scrape_saratogaspagolf_menu(url)

if __name__ == '__main__':
    main()

